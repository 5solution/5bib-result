#!/usr/bin/env python3
"""
F-024 Contract Management — DOCX template placeholder injection script.

Mục đích: Lấy file mẫu gốc Danny gửi trong `.5bib-workflow/.../templates-input/`
+ inject placeholder docxtemplater (`{varName}`) cho các hardcoded value
+ save ra `backend/assets/contract-templates/<output>.docx`.

CHIẾN LƯỢC:
  1. Mở file mẫu DOCX bằng python-docx (preserve toàn bộ format gốc — font, size,
     bold, table, image, header, footer, list, page break, etc.).
  2. Replace hardcoded text bằng placeholder qua paragraph.runs + table cells.
  3. Wrap line items table với loop marker `{#lineItems}...{/lineItems}`.
  4. Save → docxtemplater có thể render với context production.

QUAN TRỌNG:
  - Replace KHÔNG split run — dùng paragraph-level text replace giữ format.
  - Sau khi inject xong, MỞ + RE-SAVE qua LibreOffice headless (`soffice --convert-to
    docx`) để LibreOffice normalize text runs → docxtemplater resolve được placeholder
    nằm trên multiple runs.

Yêu cầu local:
  - python-docx (đã có)
  - LibreOffice cmd `soffice` (Mac: /Applications/LibreOffice.app/Contents/MacOS/soffice)
"""

import os
import sys
import shutil
import subprocess
import re
from copy import deepcopy
from docx import Document
from docx.oxml.ns import qn

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
INPUT_DIR = os.path.abspath(os.path.join(
    ROOT, "..", ".5bib-workflow", "features",
    "FEATURE-024-contract-management", "templates-input"))
# fallback to main repo path
if not os.path.exists(INPUT_DIR):
    INPUT_DIR = os.path.abspath(os.path.join(
        ROOT, "..", "..", "..", "..", ".5bib-workflow", "features",
        "FEATURE-024-contract-management", "templates-input"))
OUT_DIR = os.path.join(ROOT, "assets", "contract-templates")
TMP_DIR = "/tmp/F-024-template-build"

# ----- Bản đồ thay thế chung -----
# Mỗi entry: (text_gốc, placeholder). Apply theo thứ tự — đặt CHUỖI DÀI hơn TRƯỚC
# tránh substring collision.
COMMON_REPLACEMENTS = [
    # --- Client (Bên A — CTY CP THÀNH AN MEDIA mẫu trong gốc) ---
    ("CÔNG TY CỔ PHẦN THÀNH AN MEDIA", "{client.entityName}"),
    ("CÔNG TY CP THÀNH AN MEDIA", "{client.entityName}"),
    ("TM01-22, Tòa W1, Vinhomes West Point, đường Đỗ Đức Dục, Phường Từ Liêm, Thành phố Hà Nội",
     "{client.address}"),
    ("TM01-22, Tòa W1, Vinhomes West Point, đường Đỗ Đức Dục, Phường Từ Liêm, Hà Nội",
     "{client.address}"),
    ("0110446252", "{client.taxId}"),
    ("Vũ Phan Anh", "{client.representative}"),
    ("Tổng Giám Đốc", "{client.position}"),
    ("Tổng Giám đốc", "{client.position}"),
    # phone — client + provider trùng số trong file gốc; ưu tiên gán client
    # bởi thường client gọi trước; provider sẽ được phục hồi bằng explicit pass-2 ở
    # function inject_per_file (nếu khác).
    ("0985 737 168", "{client.phone}"),

    # --- Provider (Bên B / 5BIB / 5SOLUTION) ---
    ("CÔNG TY CỔ PHẦN 5BIB", "{provider.entityName}"),
    ("5SOLUTION TECHNOLOGY JOINT STOCK COMPANY", "{provider.entityName}"),
    ("VP905, Tầng 9 (Sàn văn phòng), Khối C, Tòa nhà Hồ Gươm Plaza, Số 102 Phố Trần Phú, Phường Hà Đông, Thành phố Hà Nội",
     "{provider.address}"),
    ("Tầng 9, Tòa nhà Hồ Gươm Plaza (tòa văn phòng), Số 102 Phố Trần Phú, Phường Hà Đông",
     "{provider.address}"),
    ("Văn phòng 501, tầng 5, tòa nhà Dreamland Bonanza",
     "{provider.address}"),
    ("0110398986", "{provider.taxId}"),
    ("Ông Nguyễn Bình Minh", "{provider.representative}"),
    ("Nguyễn Bình Minh", "{provider.representative}"),
    ("Giám đốc", "{provider.position}"),
    # Bank info
    ("110398986  - tại ngân hàng MB Bank chi nhánh Thụy Khuê",
     "{provider.bankAccount} - tại ngân hàng {provider.bankName}"),
    ("110398986 - tại ngân hàng MB Bank chi nhánh Thụy Khuê",
     "{provider.bankAccount} - tại ngân hàng {provider.bankName}"),

    # --- Race info ---
    ("Hành Trình Theo Chân Bác - Vì An Ninh Tổ Quốc", "{raceName}"),
    ("Hành Trình Theo Chân Bác – Vì An Ninh Tổ Quốc 2026", "{raceName}"),
    ("Hành Trình Theo Chân Bác – Vì An Ninh Tổ Quốc", "{raceName}"),

    # --- Contract metadata ---
    ("11.04/2026/HĐDV/TAM-5BIB", "{contractNumber}"),
    ("11.04/2026/HDDV/TAM-5BIB", "{contractNumber}"),
    ("Hôm nay, ngày    tháng    năm 2026",
     "Hôm nay, ngày {signDay} tháng {signMonth} năm {signYear}"),
    ("Năm 2026", "Năm {signYear}"),

    # --- Financial summary ---
    ("164.160.000 VND", "{totalAmount} VND"),
    ("Một trăm sáu mươi tư triệu một trăm sáu mươi ngàn đồng",
     "{totalAmountInWords}"),

    # --- Payment terms ---
    ("12%/năm", "{paymentTerms.latePenaltyRate}{paymentTerms.latePenaltyUnit}"),
    ("12% /năm", "{paymentTerms.latePenaltyRate}{paymentTerms.latePenaltyUnit}"),
    ("(mười hai phần trăm một năm)", ""),
    ("50% (năm mươi phần trăm)", "{paymentTerms.advancePercentage}%"),
    ("30 ngày làm việc",
     "{paymentTerms.paymentDeadlineDays} ngày làm việc"),
]


def replace_in_paragraph(paragraph, old, new):
    """Replace old→new in paragraph text. Cần handle multi-run text.

    Strategy: nếu paragraph.text chứa old → concat full text → replace → ghi
    về run[0] và clear các run còn lại của ĐÚNG segment. Để giữ format, ta
    chỉ replace run-by-run nếu old nằm trong 1 run; nếu nằm cross-run, fallback
    set toàn bộ paragraph text về 1 run (chấp nhận mất bold/italic CỤC BỘ tại
    chỗ đó nhưng giữ paragraph-level style).
    """
    if old not in paragraph.text:
        return False
    # Single-run fast path
    for run in paragraph.runs:
        if old in run.text:
            run.text = run.text.replace(old, new)
            return True
    # Cross-run fallback: rewrite via runs[0]
    # Tìm exact start..end run indices
    full = paragraph.text
    idx = full.find(old)
    if idx == -1:
        return False
    # Đếm cum length để biết run start/end
    cum = 0
    starts = []
    for r in paragraph.runs:
        starts.append((cum, cum + len(r.text), r))
        cum += len(r.text)
    end = idx + len(old)
    affected = [(s, e, r) for (s, e, r) in starts if not (e <= idx or s >= end)]
    if not affected:
        return False
    # Build new text for first affected run, clear rest
    first_s, first_e, first_r = affected[0]
    last_s, last_e, last_r = affected[-1]
    prefix = first_r.text[: idx - first_s]
    suffix = last_r.text[end - last_s :]
    first_r.text = prefix + new + suffix
    # Clear middle runs
    for (s, e, r) in affected[1:]:
        r.text = ""
    return True


def apply_replacements(doc, replacements, provider_overrides=None):
    """Apply COMMON_REPLACEMENTS + provider_overrides to all paragraphs + tables.

    provider_overrides: dict {old_text: new_placeholder} áp dụng SAU common
    để override (e.g., provider phone khác client phone).
    """
    all_replacements = list(replacements)
    if provider_overrides:
        all_replacements.extend(provider_overrides.items())
    # Sort by length DESC to avoid substring conflict
    all_replacements.sort(key=lambda x: -len(x[0]))

    def walk_paragraphs(paragraphs):
        for p in paragraphs:
            for old, new in all_replacements:
                replace_in_paragraph(p, old, new)

    def walk_tables(tables):
        for t in tables:
            for row in t.rows:
                for cell in row.cells:
                    walk_paragraphs(cell.paragraphs)
                    if cell.tables:
                        walk_tables(cell.tables)

    walk_paragraphs(doc.paragraphs)
    walk_tables(doc.tables)
    # Headers & footers
    for section in doc.sections:
        for hf in (section.header, section.footer,
                   section.even_page_header, section.even_page_footer):
            if hf is not None:
                walk_paragraphs(hf.paragraphs)
                walk_tables(hf.tables)


def wrap_line_items_loop(doc, table_index, header_row_count=1, columns=None):
    """Wrap loop marker quanh data rows của line-items table.

    Approach (docxtemplater table-row loop):
      - Đặt `{#lineItems}` vào CELL ĐẦU TIÊN của data row template
        + `{/lineItems}` vào CELL CUỐI CÙNG của CÙNG row đó.
      - docxtemplater nhận diện loop nằm trong 1 `<w:tr>` → duplicate row đó
        cho từng item trong array.
      - Mỗi data cell có placeholder tương ứng (`{stt}`, `{description}`, ...)
      - Xóa các data row khác (chỉ giữ 1 template row + group header).

    columns: list placeholder cho mỗi cột. Mặc định cho schema [#, Mục, Mô tả,
    Đơn giá, Số lượng, CK, Thành tiền] (Timing/RACEKIT/TICKET_SALES).
    Operations dùng [#, Hạng mục, ĐVT, Số lượng, Đơn giá, CK, Thành tiền]
    → caller override.
    """
    if table_index >= len(doc.tables):
        return
    tbl = doc.tables[table_index]
    if columns is None:
        columns = ['{stt}', '{description}', '{description}', '{unitPrice}',
                   '{quantity}', '{discount}', '{amount}']

    # Find first DATA row (skip header_row_count + group rows nếu có)
    data_rows = []
    for ri, row in enumerate(tbl.rows):
        if ri < header_row_count:
            continue
        cell_texts = [c.text.strip() for c in row.cells]
        if len(set(cell_texts)) == 1 and cell_texts[0]:
            continue
        data_rows.append(row)
    if not data_rows:
        return
    template_row = data_rows[0]

    # Set placeholder cho TEMPLATE ROW
    cells = template_row.cells
    for ci, cell in enumerate(cells):
        ph = columns[ci] if ci < len(columns) else ''
        first_p = cell.paragraphs[0]
        for r in list(first_p.runs):
            r.text = ''
        if first_p.runs:
            first_p.runs[0].text = ph
        else:
            first_p.add_run(ph)
        for extra in cell.paragraphs[1:]:
            extra._element.getparent().remove(extra._element)

    # Wrap: prepend `{#lineItems}` vào first cell + append `{/lineItems}` vào last cell
    first_cell = cells[0]
    fp = first_cell.paragraphs[0]
    # Prepend by setting first run's text = `{#lineItems}` + existing
    if fp.runs:
        fp.runs[0].text = '{#lineItems}' + fp.runs[0].text
    else:
        fp.add_run('{#lineItems}' + columns[0])

    last_cell = cells[-1]
    lp = last_cell.paragraphs[0]
    if lp.runs:
        lp.runs[-1].text = lp.runs[-1].text + '{/lineItems}'
    else:
        lp.add_run(columns[-1] + '{/lineItems}')

    # Xóa các data row khác (sau template row)
    for row in data_rows[1:]:
        row._element.getparent().remove(row._element)


def libreoffice_normalize(src_path, dst_path):
    """Normalize DOCX so docxtemplater can resolve placeholders.

    Strategy:
      1) If LibreOffice (`soffice`) is available locally → use it (best result
         because LO re-renders everything and merges adjacent runs naturally).
      2) Else → in-process XML pass: re-zip docx with text-run merge done via
         lxml. Adjacent <w:r> siblings inside the same <w:p> are merged when
         they share identical <w:rPr> (formatting). This is enough to make
         placeholders like `{varName}` end up in a SINGLE run.
    """
    soffice = "/Applications/LibreOffice.app/Contents/MacOS/soffice"
    if not os.path.exists(soffice):
        for candidate in ("/usr/bin/soffice", "/usr/local/bin/soffice", "soffice"):
            if shutil.which(candidate):
                soffice = candidate
                break
    if os.path.exists(soffice) or shutil.which(soffice):
        tmp = os.path.join(TMP_DIR, "normalize")
        os.makedirs(tmp, exist_ok=True)
        subprocess.run(
            [soffice, "--headless", "--convert-to", "docx", "--outdir", tmp, src_path],
            check=True, capture_output=True,
        )
        base = os.path.basename(src_path)
        normalized = os.path.join(tmp, base)
        if not os.path.exists(normalized):
            cands = [f for f in os.listdir(tmp) if f.endswith('.docx')]
            if cands:
                normalized = os.path.join(tmp, cands[0])
        shutil.copy(normalized, dst_path)
        return

    # In-process fallback — merge adjacent runs with identical formatting +
    # additionally repair split placeholders (`{varName}` that crosses run
    # boundaries) by stitching them into the FIRST run of each affected
    # paragraph.
    print("  (using in-process XML normalize — LibreOffice not installed)")
    import zipfile
    from lxml import etree

    W_NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    nsmap = {'w': W_NS}

    def rpr_signature(rpr):
        if rpr is None:
            return ''
        return etree.tostring(rpr, method='c14n').decode()

    def merge_runs_in_paragraph(p):
        runs = p.findall(f'{{{W_NS}}}r')
        i = 0
        while i < len(runs) - 1:
            a, b = runs[i], runs[i + 1]
            # Don't merge if either contains break / drawing / fld / etc.
            a_special = [c for c in a if c.tag not in
                         (f'{{{W_NS}}}rPr', f'{{{W_NS}}}t')]
            b_special = [c for c in b if c.tag not in
                         (f'{{{W_NS}}}rPr', f'{{{W_NS}}}t')]
            if a_special or b_special:
                i += 1
                continue
            a_rpr = a.find(f'{{{W_NS}}}rPr')
            b_rpr = b.find(f'{{{W_NS}}}rPr')
            if rpr_signature(a_rpr) != rpr_signature(b_rpr):
                i += 1
                continue
            a_t = a.find(f'{{{W_NS}}}t')
            b_t = b.find(f'{{{W_NS}}}t')
            if a_t is None or b_t is None:
                i += 1
                continue
            a_t.text = (a_t.text or '') + (b_t.text or '')
            a_t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
            b.getparent().remove(b)
            runs = p.findall(f'{{{W_NS}}}r')
        return p

    def repair_split_placeholders(p):
        """Stitch `{...}` placeholders cross-run into the first affected run.

        Approach: build full paragraph text + run boundaries. For each '{' →
        '}' pair, if start/end fall in different runs, MOVE the text from
        boundary runs so the entire `{...}` ends up in run[start_run].
        """
        runs = p.findall(f'{{{W_NS}}}r')
        if not runs:
            return
        run_texts = []
        for r in runs:
            t = r.find(f'{{{W_NS}}}t')
            run_texts.append((r, t, t.text or '' if t is not None else ''))
        full = ''.join(rt[2] for rt in run_texts)
        if '{' not in full or '}' not in full:
            return
        # Find balanced { } pairs
        i = 0
        repaired = False
        while True:
            start = full.find('{', i)
            if start == -1:
                break
            end = full.find('}', start + 1)
            if end == -1:
                break
            # Map start/end → run index
            cum = 0
            start_run = end_run = -1
            for ri, (_, _, txt) in enumerate(run_texts):
                if start_run == -1 and start < cum + len(txt):
                    start_run = ri
                    start_off = start - cum
                if end < cum + len(txt):
                    end_run = ri
                    end_off = end - cum
                    break
                cum += len(txt)
            if start_run != -1 and end_run != -1 and start_run != end_run:
                # Stitch: take full {...} string and put into run[start_run].
                # Remove from later runs.
                placeholder_str = full[start : end + 1]
                # Run start: prefix + placeholder + (suffix from end_run after end_off+1)
                _, st, stx = run_texts[start_run]
                prefix = stx[:start_off]
                _, et, etx = run_texts[end_run]
                suffix_after = etx[end_off + 1 :]
                new_start_text = prefix + placeholder_str + suffix_after
                if st is None:
                    # create <w:t>
                    new_t = etree.SubElement(run_texts[start_run][0],
                                             f'{{{W_NS}}}t')
                    new_t.text = new_start_text
                    new_t.set('{http://www.w3.org/XML/1998/namespace}space',
                              'preserve')
                else:
                    st.text = new_start_text
                    st.set('{http://www.w3.org/XML/1998/namespace}space',
                           'preserve')
                # Clear runs between start_run+1 .. end_run inclusive
                for ri in range(start_run + 1, end_run + 1):
                    _, t, _ = run_texts[ri]
                    if t is not None:
                        t.text = ''
                # Rebuild full + run_texts
                run_texts = []
                for r in runs:
                    t = r.find(f'{{{W_NS}}}t')
                    run_texts.append((r, t, t.text or '' if t is not None else ''))
                full = ''.join(rt[2] for rt in run_texts)
                repaired = True
                i = end + 1
            else:
                i = end + 1
        return repaired

    with zipfile.ZipFile(src_path, 'r') as zin:
        names = zin.namelist()
        os.makedirs(TMP_DIR, exist_ok=True)
        out_path = os.path.join(TMP_DIR, "normalized.docx")
        with zipfile.ZipFile(out_path, 'w', zipfile.ZIP_DEFLATED) as zout:
            for n in names:
                data = zin.read(n)
                if n.endswith('.xml') and (
                    n.startswith('word/document')
                    or n.startswith('word/header')
                    or n.startswith('word/footer')
                ):
                    root = etree.fromstring(data)
                    for p in root.iter(f'{{{W_NS}}}p'):
                        # Repair BEFORE merging so we don't merge across
                        # placeholder boundaries first
                        repair_split_placeholders(p)
                        merge_runs_in_paragraph(p)
                    data = etree.tostring(
                        root, xml_declaration=True, encoding='UTF-8',
                        standalone=True,
                    )
                zout.writestr(n, data)
    shutil.copy(out_path, dst_path)


# Column placeholder schemas
COLS_DEFAULT = ['{stt}', '{description}', '{description}', '{unitPrice}',
                '{quantity}', '{discount}', '{amount}']
# Operations & acceptance-operations: [#, Hạng mục, ĐVT, Số lượng, Đơn giá, ...]
COLS_OPERATIONS = ['{stt}', '{description}', '{unit}', '{quantity}',
                   '{unitPrice}', '{discount}', '{amount}', '{note}']
# Ticket sales: STT, QUYỀN LỢI, ĐVT, Số lượng (4 cols)
COLS_TICKET_SALES = ['{stt}', '{description}', '{unit}', '{quantity}']
# Acceptance Timing/Racekit: 8 cols (giống contract Timing + extra "actual"
# column → placeholder cuối là `{actualAmount}`)
COLS_ACCEPTANCE_TIMING = ['{stt}', '{description}', '{description}',
                          '{unitPrice}', '{quantity}', '{discount}',
                          '{amount}', '{actualAmount}']

# --- File-specific configs ---
TEMPLATE_CONFIGS = [
    {
        "input": "[Timing] - 5BIB - Hợp đồng dịch vụ tính giờ Hành Trình Theo Chân Bác - Vì An Ninh Tổ Quốc.docx",
        "output": "contract-timing.docx",
        "line_items_table_index": 5,
        "line_items_header_rows": 2,  # row 0 = column header, row 1 = group merged header
        "columns": COLS_DEFAULT,
    },
    {
        "input": "[RACEKIT] - 5BIB - Hợp đồng vận hành racekit - Hành Trình Theo Chân Bác - Vì An Ninh Tổ Quốc.docx",
        "output": "contract-racekit.docx",
        "line_items_table_index": 4,
        "line_items_header_rows": 1,
        "columns": COLS_DEFAULT,
    },
    {
        "input": "14.4.26 [Hành Trình vì an ninh tổ quốc] - 5Sport - Hợp đồng vận hành (1).docx",
        "output": "contract-operations.docx",
        "line_items_table_index": 5,
        "line_items_header_rows": 1,
        "columns": COLS_OPERATIONS,
    },
    {
        "input": "[5BIB] Hợp đồng bán vé Giải chạy Hành Trình Theo Chân Bác – Vì An Ninh Tổ Quốc 2026.docx",
        "output": "contract-ticket-sales.docx",
        "line_items_table_index": 7,
        "line_items_header_rows": 1,
        "columns": COLS_TICKET_SALES,
    },
    {
        "input": "[5BIB-Timing] Biên bản nghiệm thu Hành Trình Theo Chân Bác - Vì An Ninh Tổ Quốc.docx",
        "output": "acceptance-timing.docx",
        "line_items_table_index": 0,
        "line_items_header_rows": 2,
        "columns": COLS_ACCEPTANCE_TIMING,
    },
    {
        "input": "[5BIB-Racekit] Biên bản nghiệm thu Hành Trình Theo Chân Bác - Vì An Ninh Tổ Quốc.docx",
        "output": "acceptance-racekit.docx",
        "line_items_table_index": 0,
        "line_items_header_rows": 2,
        "columns": COLS_ACCEPTANCE_TIMING,
    },
    {
        "input": "[5Sport - Vận hành] Biên bản nghiệm thu Hành Trình Theo Chân Bác - Vì An Ninh Tổ Quốc.docx",
        "output": "acceptance-operations.docx",
        "line_items_table_index": 0,
        "line_items_header_rows": 1,
        "columns": COLS_OPERATIONS,
    },
]


def process_one(cfg):
    src = os.path.join(INPUT_DIR, cfg["input"])
    if not os.path.exists(src):
        print(f"✗ Input not found: {src}")
        return False
    print(f"→ Processing {cfg['output']}")
    os.makedirs(TMP_DIR, exist_ok=True)
    intermediate = os.path.join(TMP_DIR, f"INJECTED-{cfg['output']}")
    final = os.path.join(OUT_DIR, cfg["output"])

    doc = Document(src)
    apply_replacements(doc, COMMON_REPLACEMENTS)
    try:
        wrap_line_items_loop(
            doc, cfg["line_items_table_index"],
            cfg["line_items_header_rows"],
            columns=cfg.get("columns"),
        )
    except Exception as e:
        print(f"  ⚠️  Could not wrap line-items loop: {e}")
    doc.save(intermediate)
    libreoffice_normalize(intermediate, final)
    size = os.path.getsize(final)
    print(f"  ✓ Saved {final} ({size:,} bytes)")
    return True


def main():
    if not os.path.exists(INPUT_DIR):
        print(f"✗ Input dir not found: {INPUT_DIR}")
        sys.exit(1)
    os.makedirs(OUT_DIR, exist_ok=True)
    ok = 0
    for cfg in TEMPLATE_CONFIGS:
        if process_one(cfg):
            ok += 1
    print(f"\nDone: {ok}/{len(TEMPLATE_CONFIGS)} templates built")


if __name__ == "__main__":
    main()
